#!/usr/bin/env python3
"""Create the 2026-08-02 supplemental organizational resolution for 804 Quitman."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ENTITY = "Lion LFTY0435 DAO LLC"
PROPERTY = "804 S. Quitman St., Denver, CO 80219"
EFFECTIVE_DATE = "August 2, 2026"
CAPITAL_ACCOUNT_CHANGES = [
    ("Nathaniel Gipson", "Jun.–Aug. 2026", "+$8,100.00", "$210,398.29"),
    ("NARWALL Holdings, LLC", EFFECTIVE_DATE, "Restored to current schedule", "$22,000.00"),
    ("EVCO Holdings, LLC", EFFECTIVE_DATE, "Restored to current schedule", "$14,000.00"),
    ("Daniel Murrey", "Jan. 9, 2026", "Admitted; +$10,000.00", "$10,000.00"),
]
TOTAL_MEMBER_CAPITAL = "$448,398.29"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E78")
    borders.append(bottom)
    p_pr.append(borders)


def add_resolution_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.first_line_indent = Inches(0.25)
    lead, rest = text.split(",", 1)
    run = p.add_run(lead + ",")
    run.bold = True
    p.add_run(rest)


def add_signature_block(
    doc: Document,
    name: str,
    title: str | None = None,
    signature_image: Path | None = None,
    signed_date: str | None = None,
) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.9)
    table.columns[1].width = Inches(1.45)
    table.cell(0, 0).width = Inches(4.9)
    table.cell(0, 1).width = Inches(1.45)
    if signature_image:
        signature_cell = table.cell(0, 0)
        signature_cell.text = ""
        signature_run = signature_cell.paragraphs[0].add_run()
        signature_run.add_picture(str(signature_image), width=Inches(2.8))
    else:
        table.cell(0, 0).text = "____________________________________________"
    table.cell(0, 1).text = signed_date or "________________"
    table.cell(1, 0).text = name if not title else f"{name}\n{title}"
    table.cell(1, 1).text = "Date"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_cell_margins(cell, top=0, bottom=0)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.keep_together = True
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def create_document(output: Path, manager_signature_image: Path | None = None) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(33, 37, 41)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("SUPPLEMENTAL ORGANIZATIONAL RESOLUTION")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 78, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(1)
    run = subtitle.add_run(f"OF THE MEMBERS OF {ENTITY}")
    run.bold = True
    run.font.size = Pt(11.5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run(f"Effective {EFFECTIVE_DATE}  •  {PROPERTY}")
    r.italic = True
    r.font.color.rgb = RGBColor(89, 89, 89)
    add_rule(doc)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(9)
    intro.paragraph_format.line_spacing = 1.08
    intro.add_run(
        f"The undersigned members of {ENTITY}, a Wyoming limited liability "
        "company (the \"Company\"), act by written consent under the Operating "
        "Agreement dated July 23, 2025 (the \"Operating Agreement\"). This Supplemental "
        "Organizational Resolution amends and supplements the organizational resolution included "
        "with that Operating Agreement. All provisions not expressly modified below remain in force."
    )

    add_resolution_paragraph(
        doc,
        "RESOLVED, that EVCO Holdings, LLC and NARWALL Holdings, LLC "
        "are confirmed and restored as current members and parties to the Operating Agreement. "
        "Earl V. Co is authorized to execute this Resolution as Manager of EVCO Holdings, LLC and "
        "NARWALL Holdings, LLC.",
    )
    add_resolution_paragraph(
        doc,
        "FURTHER RESOLVED, that Daniel Murrey is admitted as a member of the Company in recognition "
        "of his documented $10,000.00 equity contribution recorded on January 9, 2026.",
    )
    add_resolution_paragraph(
        doc,
        "FURTHER RESOLVED, that the Company ratifies the following additional capital-account "
        "credits to Nathaniel Gipson: $2,700.00 for June 2026, $2,700.00 for July 2026, and "
        "$2,700.00 accrued for August 2026, producing a recorded capital-account balance of "
        "$210,398.29 through August 2026.",
    )
    add_resolution_paragraph(
        doc,
        "FURTHER RESOLVED, that the additional contributions and credits described above are "
        "treated as equity contributions at the applicable premium. Their final units, ownership "
        "percentage, and profit-and-loss allocation will be entered on the Company cap table after "
        "the premium conversion is applied consistently with the Operating Agreement and the "
        "Company's valuation methodology. No ownership percentage is created or implied solely by "
        "dividing these dollar amounts by total recorded capital.",
    )
    add_resolution_paragraph(
        doc,
        f"FURTHER RESOLVED, that after recording the changes below, total current outstanding member "
        f"and equity capital accounts as of {EFFECTIVE_DATE} equal {TOTAL_MEMBER_CAPITAL}. Unchanged "
        "balances remain as previously recorded. This amount is an "
        "accounting balance and does not itself establish ownership percentages or the cash "
        "currently held by the Company.",
    )
    add_resolution_paragraph(
        doc,
        "FURTHER RESOLVED, that, solely as required or reasonably necessary to complete an "
        "authorized refinance of the Property, record title to the Property may be vested in "
        "Nathaniel Gipson. Any such vesting is nominee and financing accommodation only and does "
        "not convey, alter, or extinguish the Company's beneficial or economic ownership of the "
        "Property. Beneficial ownership, member interests, allocations, distributions, and proceeds "
        "remain governed by the Operating Agreement, this Resolution, the Company's capital "
        "accounts, and its final cap table. Nathaniel Gipson shall hold record title subject to those "
        "rights and shall execute any deed, trust instrument, or other document reasonably required "
        "to preserve or restore the Company's beneficial ownership after the refinance.",
    )
    add_resolution_paragraph(
        doc,
        "FURTHER RESOLVED, that the Manager is authorized to update the Company's membership "
        "ledger, capital accounts, cap table, tax records, and related schedules to implement this "
        "Resolution, provided those updates preserve the premium treatment stated above.",
    )
    add_resolution_paragraph(
        doc,
        "FURTHER RESOLVED, that this Resolution may be executed electronically and in "
        "counterparts, each of which is deemed an original and all of which together form one instrument.",
    )

    summary_heading = doc.add_paragraph()
    summary_heading.paragraph_format.space_before = Pt(3)
    summary_heading.paragraph_format.space_after = Pt(5)
    keep_with_next(summary_heading)
    r = summary_heading.add_run("Changes since the prior organizational resolution")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 120)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Member", "Effective period", "Change", "Resulting capital"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, "1F4E78")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for idx, values in enumerate(CAPITAL_ACCOUNT_CHANGES, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            if idx % 2 == 0:
                set_cell_shading(cells[i], "EAF2F8")
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8.5)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(8)
    note.add_run("Debt status: ").bold = True
    note.add_run(
        "Mortgage and other third-party loan balances are maintained separately under their "
        "applicable amortization and payoff records. The prior Mt. Vernon DAO loan has been paid "
        "in full and is not an outstanding Company liability."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        "IN WITNESS WHEREOF, the undersigned members approve this Resolution, and Daniel Murrey "
        "accepts admission as a member, effective as stated above."
    ).bold = True

    doc.add_section(WD_SECTION.NEW_PAGE)
    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_title.paragraph_format.space_after = Pt(3)
    r = sig_title.add_run("MEMBER CONSENTS AND JOINING MEMBER ACCEPTANCE")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(31, 78, 120)
    sig_sub = doc.add_paragraph()
    sig_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_sub.paragraph_format.space_after = Pt(18)
    sig_sub.add_run(f"{ENTITY} • Effective {EFFECTIVE_DATE}").italic = True

    manager_date = "08/02/2026" if manager_signature_image else None
    add_signature_block(
        doc,
        "Earl V. Co",
        "Manager, ECO Systems, LLC",
        manager_signature_image,
        manager_date,
    )
    add_signature_block(
        doc,
        "Earl V. Co",
        "Manager, EVCO Holdings, LLC",
        manager_signature_image,
        manager_date,
    )
    add_signature_block(
        doc,
        "Earl V. Co",
        "Manager, NARWALL Holdings, LLC",
        manager_signature_image,
        manager_date,
    )
    add_signature_block(doc, "Nathaniel Gipson")
    add_signature_block(doc, "Wesley Babcock")
    add_signature_block(doc, "Ian Haber")
    add_signature_block(doc, "Brandon McArthur")
    add_signature_block(doc, "Kyle Randal McArthur")
    joining = doc.add_paragraph()
    joining.paragraph_format.space_before = Pt(7)
    joining.paragraph_format.space_after = Pt(8)
    r = joining.add_run("JOINING MEMBER ACCEPTANCE")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 120)
    joining.add_run(
        " — I accept admission as a member and agree to be bound by the Operating Agreement, "
        "as supplemented by this Resolution."
    )
    add_signature_block(doc, "Daniel Murrey")

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"{ENTITY} — Supplemental Organizational Resolution — {EFFECTIVE_DATE}")
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 110, 110)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def validate(output: Path) -> None:
    doc = Document(output)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    required = [
        ENTITY,
        "Daniel Murrey",
        "$10,000.00",
        "EVCO Holdings, LLC",
        "$14,000.00",
        "NARWALL Holdings, LLC",
        "$22,000.00",
        "January 9, 2026",
        "Nathaniel Gipson",
        "$210,398.29",
        TOTAL_MEMBER_CAPITAL,
        "The prior Mt. Vernon DAO loan has been paid in full",
        "$2,700.00 for June 2026",
        "$2,700.00 for July 2026",
        "$2,700.00 accrued for August 2026",
        "premium",
        "record title to the Property may be vested in Nathaniel Gipson",
        "does not convey, alter, or extinguish the Company's beneficial or economic ownership",
        EFFECTIVE_DATE,
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Validation failed; missing: {missing}")
    forbidden = ["Solum Algo", "SolumAlgo", "52.793%", "63.56%"]
    present = [item for item in forbidden if item in text]
    if present:
        raise RuntimeError(f"Validation failed; forbidden stale/handle text present: {present}")
    if text.count("Daniel Murrey") < 3:
        raise RuntimeError("Validation failed; Daniel Murrey is not present in resolution, table, and signature block")
    for signer in ["Brandon McArthur", "Kyle Randal McArthur"]:
        if text.count(signer) != 1:
            raise RuntimeError(f"Validation failed; {signer} must appear only in a signature block")
    if "Thomas A. Austin" in text:
        raise RuntimeError("Validation failed; unchanged non-signing member should not be restated")
    if text.count("Manager, EVCO Holdings, LLC") != 1:
        raise RuntimeError("Validation failed; EVCO Holdings manager signature block is missing or duplicated")
    if text.count("Manager, NARWALL Holdings, LLC") != 1:
        raise RuntimeError("Validation failed; NARWALL Holdings manager signature block is missing or duplicated")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    parser.add_argument("--manager-signature-image", type=Path)
    args = parser.parse_args()
    if args.manager_signature_image and not args.manager_signature_image.is_file():
        raise SystemExit(f"Manager signature image not found: {args.manager_signature_image}")
    create_document(args.output, args.manager_signature_image)
    validate(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
